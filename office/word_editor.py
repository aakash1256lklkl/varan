"""
Varan Word editor — create, edit, read and summarize .docx files.

Creation/reading/summarizing are built on python-docx.
Editing routes through DocxEngine (https://github.com/ruwadgroup/docxengine),
which edits the OOXML directly so formatting, tracked changes and comments are
preserved — unlike naive run-clearing find-and-replace.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from docxengine import call as _docx_call

from . import word_live as _live


class WordEditor:
    extension = ".docx"

    # ------------------------------------------------------------------
    # Creation
    # ------------------------------------------------------------------
    def create_document(self, path: str | Path, title: str = "", body: list[Any] | None = None) -> str:
        """Create a new Word document.

        body is a list of block dicts:
          {"type": "heading", "level": 1, "text": "..."}
          {"type": "paragraph", "text": "..." , "bold": bool, "italic": bool}
          {"type": "bullet", "text": "..."}
          {"type": "table", "headers": [...], "rows": [[...], ...]}
        """
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

        if title:
            doc.add_heading(title, level=0)

        if body:
            self._apply_blocks(doc, body)

        doc.save(str(path))
        return str(path)

    def _apply_blocks(self, doc: Document, body: list[dict]) -> None:
        for block in body or []:
            if not isinstance(block, dict):
                continue
            btype = block.get("type", "paragraph")
            text = str(block.get("text", ""))
            if btype == "heading":
                level = int(block.get("level", 1) or 1)
                level = max(1, min(level, 6))
                p = doc.add_heading(text, level=level)
                self._style_run((p.runs[0] if p.runs else p.add_run("")), block)
            elif btype == "title":
                p = doc.add_heading(text, level=0)
                self._style_run((p.runs[0] if p.runs else p.add_run("")), block)
            elif btype == "bullet":
                p = doc.add_paragraph(text, style="List Bullet")
                self._style_run((p.runs[0] if p.runs else p.add_run("")), block)
            elif btype == "numbered":
                p = doc.add_paragraph(text, style="List Number")
                self._style_run((p.runs[0] if p.runs else p.add_run("")), block)
            elif btype in ("divider", "scene_break"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("* * *")
            elif btype == "table":
                self._add_table(doc, block)
            else:  # paragraph
                self._style_run(doc.add_paragraph().add_run(text), block)

    @staticmethod
    def _style_run(run, block: dict) -> None:
        """Apply optional font / bold / italic / size styling to a run from a
        content block (also accepts 'courier': true as a shorthand for a
        Courier New font, used for script/teletype content)."""
        fname = block.get("font") or (block.get("courier") and "Courier New")
        if fname:
            run.font.name = str(fname)
        if block.get("bold"):
            run.bold = True
        if block.get("italic"):
            run.italic = True
        if block.get("size"):
            run.font.size = Pt(float(block["size"]))

    def insert_content(self, path: str | Path, content: list[Any] | None,
                       inplace: bool = False) -> dict:
        """Append structured content BLOCKS to an existing .docx as real Word
        paragraphs — true Heading 1-6 styles, List Bullet / List Number, bold /
        italic runs, a Courier font, centered scene-break dividers and tables.
        This is how Varan adds a fully-styled section to an existing document
        instead of dumping plain text (or literal '##'/'-' markdown) into it.

        Returns {'path': <written file>, 'blocks': <number of blocks>}.
        """
        src = Path(path)
        if not src.exists():
            raise FileNotFoundError(src)
        doc = Document(str(src))
        self._apply_blocks(doc, list(content or []))
        dst = src if inplace else self._edited_copy(src)
        doc.save(str(dst))
        return {"path": str(dst), "blocks": len(list(content or []))}

    def _add_table(self, doc: Document, block: dict) -> None:
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        ncols = max(len(headers), *(len(r) for r in rows)) if rows else len(headers)
        nrows = (1 if headers else 0) + len(rows)
        if ncols == 0 or nrows == 0:
            return
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ri = 0
        if headers:
            for ci, h in enumerate(headers):
                cell = table.cell(0, ci)
                cell.text = str(h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            ri = 1
        for r in rows:
            vals = list(r)
            for ci in range(ncols):
                v = vals[ci] if ci < len(vals) else ""
                table.cell(ri, ci).text = str(v)
            ri += 1

    # ------------------------------------------------------------------
    # Editing (via DocxEngine — formatting-preserving)
    # ------------------------------------------------------------------
    def edit_document(self, path: str | Path, edits: list[dict] | None = None,
                      append: str | None = None, track_changes: bool = False,
                      table_edits: list[dict] | None = None,
                      author: str = "Varan", inplace: bool = False) -> str:
        """Edit an existing document with DocxEngine.

        Returns the path of the file that was written. By default a new
        "NAME_edited.docx" copy is produced; pass inplace=True to write the
        changes back over the original file (used when editing the user's
        explicitly selected target file, so no duplicate copy appears).

        edits: list of {"action": "replace"|"insert_after"|"insert_before"|"delete",
                        "match": text to find (supports search anchor or plain text),
                        "text": new text,
                        "count": N (replace: how many occurrences to replace;
                                  -1 means replace EVERY occurrence in the doc),
                        "bold": bool, "italic": bool,
                        "track": bool (per-edit override), }
        table_edits: list of {"table": <1-based table index>, "cell": "A1",
                              "text": "..."} — set the text of a cell in an
                              existing table (complex table editing).
        track_changes: write real tracked changes (redlines) by default.
        author: named author used when tracking changes.
        """
        src = Path(path)
        if not src.exists():
            raise FileNotFoundError(src)
        dst = src if inplace else self._edited_copy(src)

        opened = _docx_call("docx_open", {"path": str(src)})
        doc_id = opened["doc_id"]

        for edit in (edits or []):
            self._apply_edit(doc_id, edit, track_changes, author)

        if table_edits:
            self._apply_table_edits(doc_id, table_edits, track_changes, author)

        if append is not None:
            last_anchor = self._last_anchor(doc_id)
            if last_anchor is not None:
                _docx_call("docx_insert",
                           {"doc_id": doc_id, "after": last_anchor, "content": str(append)})

        _docx_call("docx_save", {"doc_id": doc_id, "path": str(dst)})
        return str(dst)

    # -- live editing via Word COM (document open in Microsoft Word) ------
    def is_open_in_word(self, path: str | Path) -> bool:
        """Return True if the .docx is currently open in Microsoft Word."""
        try:
            return bool(_live.is_open_in_word(path))
        except Exception:  # noqa: BLE001
            return False

    def live_edit(self, path: str | Path, edits=None, append: str | None = None,
                  page: int | None = None, remove_empty_paragraphs: bool = False,
                  content=None, table_edits=None) -> dict:
        """Apply the same edit job edit_document supports, but LIVE in the open
        Word document via COM. Returns {'ok','path','mode':'live-word'}.
        """
        return _live.live_edit(path, edits=edits, append=append, page=page,
                               remove_empty_paragraphs=remove_empty_paragraphs,
                               content=content, table_edits=table_edits)

    def delete_page(self, path: str | Path, page: int) -> dict:
        """Delete a whole numeric page from a document open in Word, live."""
        return _live.live_edit(path, page=page)

    def remove_empty_paragraphs(self, path: str | Path, inplace: bool = False) -> dict:
        """Remove every empty body paragraph from a .docx (the typical cause of
        'blank pages' — empty paragraphs, especially empty Heading 1 paragraphs
        used as page separators). Only empty paragraphs are removed; paragraphs
        that contain any text are left untouched.

        Returns {'removed': N, 'path': <written file>}. If no empty paragraphs
        are found, returns {'removed': 0, 'path': <original path>} and writes
        nothing (so a 'delete blank pages' call on an already-clean doc never
        touches the file).
        """
        src = Path(path)
        if not src.exists():
            raise FileNotFoundError(src)

        doc = Document(str(src))
        removed = 0
        for p in list(doc.paragraphs):
            if not p.text.strip():
                p._element.getparent().remove(p._element)
                removed += 1

        if removed == 0:
            return {"removed": 0, "path": str(src)}

        dst = src if inplace else self._edited_copy(src)
        doc.save(str(dst))
        return {"removed": removed, "path": str(dst)}

    def delete_range(self, path: str | Path, match: str,
                     end_match: str | None = None, end_level: int | None = None,
                     inplace: bool = False) -> dict:
        """Delete every paragraph from the first paragraph containing `match`
        through the end spanning paragraph — the LAST occurrence of `end_match`
        (inclusive), the paragraph just before the next Heading of `end_level`,
        or the document end. This is the file-based equivalent of the live Word
        delete_range: it works directly on the .docx without needing Word.

        Returns {'path': <written file>, 'deleted': N, 'from': <first text>,
        'to': <last text>}. Raises LookupError if `match` / `end_match` are not
        found — never a silent no-op.
        """
        src = Path(path)
        if not src.exists():
            raise FileNotFoundError(src)
        doc = Document(str(src))
        ps = doc.paragraphs

        start = None
        for i, p in enumerate(ps):
            if match and match in (p.text or ""):
                start = i
                break
        if start is None:
            raise LookupError(f"text {match!r} (start of delete_range) was not found")

        end = len(ps) - 1
        if end_match:
            end = None
            for i in range(len(ps) - 1, start - 1, -1):
                if end_match in (ps[i].text or ""):
                    end = i
                    break
            if end is None:
                raise LookupError(f"text {end_match!r} (end of delete_range) was not found")
        elif end_level:
            level = max(1, min(int(end_level or 1), 6))
            for i in range(start + 1, len(ps)):
                st = (ps[i].style.name if ps[i].style is not None else "") or ""
                m = re.match(r"Heading\s*([1-6])", st)
                if m and int(m.group(1)) == level:
                    end = i - 1
                    break

        first_text = (ps[start].text or "").strip()
        last_text = (ps[end].text or "").strip()
        removed = end - start + 1
        for p in ps[start:end + 1]:
            p._element.getparent().remove(p._element)

        dst = src if inplace else self._edited_copy(src)
        doc.save(str(dst))
        return {"path": str(dst), "deleted": removed,
                "from": first_text, "to": last_text}

    def get_headings(self, path: str | Path, limit: int = 100) -> list[str]:
        """Return the non-empty paragraph texts of an open Word document, so the
        model can pick a real anchor to edit/delete instead of guessing text."""
        try:
            return _live.get_headings(path, limit=limit)
        except Exception:  # noqa: BLE001
            return []

    def live_append(self, path: str | Path, text: str) -> dict:
        """Append text to a document currently open in Word, live on screen."""
        return _live.append_text(path, text)

    def _apply_edit(self, doc_id: str, edit: dict, track_changes: bool, author: str) -> None:
        action = edit.get("action", "replace")
        match = edit.get("match", "")
        text = edit.get("text", "")
        track = bool(edit.get("track", track_changes))
        kwargs = {"doc_id": doc_id, "track_changes": track, "author": author}

        if action == "delete":
            anchor = self._anchor_for(doc_id, match)
            if anchor is None:
                raise LookupError(f"text {match!r} to delete was not found")
            _docx_call("docx_delete", {**kwargs, "anchor": anchor})
            return

        if action in ("insert_after", "insert_before"):
            anchor = self._anchor_for(doc_id, match)
            if anchor is None:
                raise LookupError(f"anchor text {match!r} was not found")
            key = "after" if action == "insert_after" else "before"
            _docx_call("docx_insert", {**kwargs, key: anchor, "content": text})
            return

        # replace — keep formatting of the matched span.
        # count == -1 replaces EVERY occurrence in the whole document (complex
        # replace-all); otherwise a single anchored replacement of the first
        # match is done (default, safest when text appears once).
        count = int(edit.get("count", 1) or 1)
        if count < 0:
            _docx_call("docx_replace", {**kwargs, "old": match, "new": text, "all": True})
            return
        anchor = self._anchor_for(doc_id, match)
        if anchor is None:
            raise LookupError(f"text {match!r} to replace was not found")
        _docx_call("docx_replace", {**kwargs, "anchor": anchor, "old": match, "new": text})

    def _apply_table_edits(self, doc_id: str, table_edits: list[dict],
                           track_changes: bool, author: str) -> None:
        """Apply complex table edits: set the text of specific cells in an
        existing table, exactly like editing a spreadsheet cell but for a Word
        table. Formatting of the cell is preserved by DocxEngine."""
        outline = _docx_call("docx_outline", {"doc_id": doc_id})
        raw_tables = outline.get("tables")
        tables = list(raw_tables) if isinstance(raw_tables, list) else []
        if not tables:
            raise LookupError("no tables found in the document")
        index_of = {}
        for i, t in enumerate(tables):
            tbl = dict(t)
            num = int(tbl.get("index", i + 1) or i + 1)
            index_of[num] = str(tbl.get("anchor") or "")
        for te in (table_edits or []):
            idx = int(te.get("table", 1) or 1)
            anchor = index_of.get(idx)
            if anchor is None:
                raise LookupError(f"table {idx} was not found (document has {len(tables)})")
            cell = te.get("cell") or {}
            pts = {"anchor": anchor, "cells": [{**cell, "text": str(te.get("text", ""))}],
                   "track_changes": bool(te.get("track", track_changes)),
                   "author": author}
            _docx_call("docx_table", {"doc_id": doc_id, "op": "set_cells", **pts})

    def _anchor_for(self, doc_id: str, match: str) -> str | None:
        """Return the anchor of the first paragraph containing `match`,
        or accept a direct anchor if `match` looks like 'P2#abc'."""
        if not match:
            return None
        if re.fullmatch(r"[A-Z]\d+#[0-9a-f]+", match):
            return match
        res = _docx_call("docx_search", {"doc_id": doc_id, "query": match})
        matches = res.get("matches") or []
        if not matches:
            return None
        return matches[0]["anchor"]

    def _last_anchor(self, doc_id: str) -> str | None:
        """Return the anchor of the last body paragraph, or None if empty."""
        res = _docx_call("docx_read", {"doc_id": doc_id, "scope": "body"})
        content = res.get("content") or ""
        anchors = re.findall(r"\[([A-Z]\d+#[0-9a-f]+)\]", content)
        return anchors[-1] if anchors else None

    def _edited_copy(self, src: Path) -> Path:
        stem = src.stem
        return src.with_name(f"{stem}_edited{self.extension}")

    # ------------------------------------------------------------------
    # Reading
    # ------------------------------------------------------------------
    def read_document(self, path: str | Path) -> dict:
        src = Path(path)
        if not src.exists():
            raise FileNotFoundError(src)
        doc = Document(str(src))
        paragraphs = [{"text": p.text, "style": p.style.name} for p in doc.paragraphs if p.text.strip()]
        tables = []
        for t in doc.tables:
            rows = []
            for row in t.rows:
                rows.append([c.text for c in row.cells])
            tables.append(rows)
        return {"path": str(src), "paragraphs": paragraphs, "tables": tables}

    def summarize(self, path: str | Path, max_points: int = 10) -> dict:
        data = self.read_document(path)
        headings = [p["text"] for p in data["paragraphs"] if p["style"].startswith("Heading")]
        full_text = "\n".join(p["text"] for p in data["paragraphs"])
        paragraphs = [p["text"] for p in data["paragraphs"] if not p["style"].startswith("Heading")]
        return {
            "title": headings[0] if headings else Path(path).stem,
            "headings": headings,
            "paragraph_count": len(paragraphs),
            "table_count": len(data["tables"]),
            "word_count": len(re.findall(r"\b\w+\b", full_text)),
            "text_preview": " ".join(paragraphs)[:2000],
        }
